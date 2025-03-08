from fastapi import FastAPI, Request, HTTPException

from fastapi import UploadFile, File
from fastapi.responses import JSONResponse
from http import HTTPStatus
from docx import Document
from io import BytesIO

from contextlib import asynccontextmanager
from psycopg_pool import AsyncConnectionPool

from fastapi import Depends
from sqlalchemy.ext.asyncio import AsyncSession
from database import get_db
from models import Test2


def get_conn_str():

    # Database configuration
    DATABASE_CONFIG = {
        "host": "pg_db",
        "port": 5432,
        "dbname": "py_docx_fastapi",
        "user": "py_docx_fastapi",
        "password": "py_docx_fastapi"
    }

    print('db_name', DATABASE_CONFIG.get('dbname'))
    return f"""
    dbname={DATABASE_CONFIG.get('dbname')}
    user={DATABASE_CONFIG.get('user')}
    password={DATABASE_CONFIG.get('password')}
    host={DATABASE_CONFIG.get('host')}
    port={DATABASE_CONFIG.get('port')}
    """


@asynccontextmanager
async def lifespan(app: FastAPI):
    app.async_pool = AsyncConnectionPool(conninfo=get_conn_str())

    # Create table if it doesn't exist
    async with app.async_pool.connection() as conn:
        async with conn.cursor() as cur:
            await cur.execute("""
                CREATE TABLE IF NOT EXISTS test_table (
                    id SERIAL PRIMARY KEY,
                    name VARCHAR(100)
                );
            """)

    yield
    await app.async_pool.close()


app = FastAPI(docs_url="/docs", lifespan=lifespan)


@app.get("/test")
async def read_test():
    return {"message": "This is a test endpoint"}


@app.post('/upload_docx', responses={HTTPStatus.BAD_REQUEST: {"model": dict}})
async def upload_docx(doc_file: UploadFile = File(...)) -> dict:
    try:
        # Read the uploaded file as bytes
        contents = await doc_file.read()

        # Load the document using python-docx
        doc = Document(BytesIO(contents))

        # Print the first 20 lines
        for i in range(min(20, len(doc.paragraphs))):
            print(doc.paragraphs[i].text)

        print("type(doc):", type(doc))

        return JSONResponse(
            status_code=HTTPStatus.CREATED,
            content={"success": True},
        )
    except Exception as e:
        return JSONResponse(
            status_code=HTTPStatus.BAD_REQUEST,
            content={"error": str(e)},
        )


@app.on_event("startup")
async def startup_event():
    pass


@app.on_event("shutdown")
async def shutdown_event():
    pass


@app.post("/test_table")
async def insert_into_test_table(name: str):
    try:
        async with app.async_pool.connection() as conn:
            async with conn.cursor() as cur:
                await cur.execute("""
                    INSERT INTO test_table (name)
                    VALUES (%s);
                """, (name,))
                await conn.commit()
                return {"message": "Inserted successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/test_table")
async def read_all_from_test_table():
    try:
        async with app.async_pool.connection() as conn:
            async with conn.cursor() as cur:
                await cur.execute("SELECT * FROM test_table;")
                results = await cur.fetchall()
                return [{"id": row[0], "name": row[1]} for row in results]
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/test2/")
async def create_test2(name: str, db: AsyncSession = Depends(get_db)):
    new_test2 = Test2(name=name)
    db.add(new_test2)
    await db.commit()
    await db.refresh(new_test2)
    return new_test2


@app.get("/test2/{test2_id}")
async def read_test2(test2_id: int, db: AsyncSession = Depends(get_db)):
    test2 = await db.get(Test2, test2_id)
    return test2 if test2 else {"error": "Test2 not found"}
